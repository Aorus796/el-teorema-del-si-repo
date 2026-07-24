#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "GDD.md"
OUTPUT = ROOT / "docs" / "El_Teorema_del_Si_GDD.docx"

ACCENT = "1F5964"
ACCENT_2 = "9A6138"
LIGHT = "EAF4F2"
LIGHT_2 = "F6F2EB"
TEXT = "263238"
MUTED = "5D6B70"
WHITE = "FFFFFF"
BORDER = "B7C8CC"
FONT = "Liberation Sans"
MONO = "Liberation Mono"


def set_run_font(run, name: str, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run("Página ")
    set_run_font(run, FONT, 8.5, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_paragraph_border(paragraph, left_color=ACCENT, left_size=18):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(left_size))
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), left_color)
    p_bdr.append(left)


def add_inline(paragraph, text: str, base_size: float = 10.5, base_color: str = TEXT, mono: bool = False):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, MONO if mono else FONT, base_size, base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, FONT, base_size, base_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, MONO, base_size - 0.5, ACCENT)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, FONT, base_size, base_color, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, MONO if mono else FONT, base_size, base_color)


def setup_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Title", 30, ACCENT, 0, 8),
        ("Subtitle", 16, MUTED, 0, 16),
        ("Heading 1", 19, ACCENT, 18, 8),
        ("Heading 2", 14, ACCENT_2, 13, 5),
        ("Heading 3", 11.5, ACCENT, 10, 3),
        ("Heading 4", 10.5, MUTED, 8, 2),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["Caption"].font.name = FONT
    styles["Caption"]._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    styles["Caption"]._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    styles["Caption"].font.size = Pt(8.5)
    styles["Caption"].font.italic = True
    styles["Caption"].font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc: Document):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EL TEOREMA DEL SÍ")
    set_run_font(r, FONT, 31, ACCENT, bold=True)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Game Design Document")
    set_run_font(r, FONT, 17, MUTED)
    p.paragraph_format.space_after = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Aventura narrativa de puzles matemáticos en pixel art")
    set_run_font(r, FONT, 11.5, TEXT, italic=True)

    for _ in range(5):
        doc.add_paragraph()

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=ACCENT, size=8)
    labels = ["Versión", "Estado", "Fecha"]
    values = ["0.1", "Diseño base aprobado provisionalmente", "24 de julio de 2026"]
    for i, (label, value) in enumerate(zip(labels, values)):
        table.cell(i, 0).width = Inches(1.35)
        table.cell(i, 1).width = Inches(4.75)
        set_cell_shading(table.cell(i, 0), ACCENT)
        set_cell_shading(table.cell(i, 1), LIGHT)
        for cell in table.rows[i].cells:
            set_cell_margins(cell, 130, 150, 130, 150)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = table.cell(i, 0).paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(label)
        set_run_font(r0, FONT, 9.5, WHITE, bold=True)
        p1 = table.cell(i, 1).paragraphs[0]
        r1 = p1.add_run(value)
        set_run_font(r1, FONT, 9.5, TEXT)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento fuente mantenido en Markdown dentro del repositorio del proyecto")
    set_run_font(r, FONT, 8.5, MUTED)
    doc.add_page_break()


def add_markdown_table(doc: Document, rows: list[list[str]]):
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:] if all(re.fullmatch(r"\s*:?-+:?\s*", c) for c in rows[1]) else rows[1:]
    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    all_rows = [header] + body
    max_lengths = []
    for c in range(cols):
        max_lengths.append(max(4, max(len(r[c]) if c < len(r) else 0 for r in all_rows)))
    weights = [min(max(x, 8), 42) for x in max_lengths]
    total = sum(weights)
    page_width = 6.8
    widths = [max(0.7, page_width * w / total) for w in weights]
    scale = page_width / sum(widths)
    widths = [w * scale for w in widths]

    prevent_row_split(table.rows[0])
    for c, text in enumerate(header):
        cell = table.rows[0].cells[c]
        cell.width = Inches(widths[c])
        set_cell_shading(cell, ACCENT)
        set_cell_margins(cell, 100, 100, 100, 100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 24 else WD_ALIGN_PARAGRAPH.LEFT
        add_inline(p, text.strip(), 8.8, WHITE)
        for run in p.runs:
            run.bold = True
    if len(body) >= 8:
        set_repeat_table_header(table.rows[0])

    for row_data in body:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for c in range(cols):
            text = row_data[c].strip() if c < len(row_data) else ""
            cell = cells[c]
            cell.width = Inches(widths[c])
            set_cell_margins(cell, 95, 100, 95, 100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            if re.fullmatch(r"[\d.,%+\-–— ]+", text) or len(text) <= 5:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, text, 8.7, TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.65)
    set_cell_shading(table.cell(0, 0), "F3F5F5")
    set_table_borders(table, color="CBD4D6", size=6)
    cell = table.cell(0, 0)
    set_cell_margins(cell, 120, 140, 120, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, MONO, 8.3, TEXT)
        if idx < len(lines) - 1:
            run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_quote(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    set_cell_shading(table.cell(0, 0), LIGHT)
    set_table_borders(table, color=ACCENT, size=8)
    cell = table.cell(0, 0)
    set_cell_margins(cell, 130, 170, 130, 170)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline(p, text, 10.2, TEXT)
    for run in p.runs:
        run.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc: Document, alt: str, path_text: str):
    path = (SOURCE.parent / path_text).resolve()
    if not path.exists():
        p = doc.add_paragraph()
        add_inline(p, f"[Imagen no encontrada: {path_text}]", 9, "B00020")
        return
    with Image.open(path) as img:
        px_w, px_h = img.size
    max_w = 6.6
    max_h = 7.2
    ratio = px_h / px_w
    width = min(max_w, max_h / ratio)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(alt)


def parse_markdown(doc: Document, text: str):
    lines = text.splitlines()
    i = 0
    first_h1 = True
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if i == 0 and stripped == "---":
            i += 1
            while i < len(lines) and lines[i].strip() != "---":
                i += 1
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, code)
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            add_markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2).strip()
            p = doc.add_paragraph(style=f"Heading {level}")
            if level == 1:
                if not first_h1 and (re.match(r"^\d+\.", title) or title in {"Apéndices", "Índice de capítulos"}):
                    p.paragraph_format.page_break_before = True
                first_h1 = False
            add_inline(p, title, {1: 19, 2: 14, 3: 11.5, 4: 10.5}[level], ACCENT if level in (1, 3) else ACCENT_2)
            for run in p.runs:
                run.bold = True
            i += 1
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            add_quote(doc, " ".join(quote_lines))
            continue
        bullet = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        numbered = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if bullet or numbered:
            m = bullet or numbered
            level = min(3, len(m.group(1)) // 2)
            p = doc.add_paragraph(style="List Bullet" if bullet else "Normal")
            p.paragraph_format.left_indent = Inches(0.2 + 0.22 * level)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            p.paragraph_format.space_after = Pt(2)
            if numbered:
                prefix = f"{m.group(2)}. "
                run = p.add_run(prefix)
                set_run_font(run, FONT, 10.2, TEXT, bold=False)
                add_inline(p, m.group(3).strip(), 10.2, TEXT)
            else:
                add_inline(p, m.group(2).strip(), 10.2, TEXT)
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith("|") or nxt.startswith(">") or nxt.startswith("```") or re.match(r"!\[.*\]\(.*\)", nxt) or re.match(r"^\s*[-*]\s+", lines[i]) or re.match(r"^\s*\d+\.\s+", lines[i]) or nxt == "---":
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines), 10.5, TEXT)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = True

    setup_styles(doc)
    add_cover(doc)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("EL TEOREMA DEL SÍ  ·  GDD 0.1")
    set_run_font(r, FONT, 8.2, MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)

    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))

    core = doc.core_properties
    core.title = "El Teorema del Sí - Game Design Document"
    core.subject = "Diseño integral de aventura narrativa de puzles matemáticos"
    core.keywords = "game design document, videojuegos, puzles, pixel art, boda"
    core.comments = "Documento consolidado del proyecto"
    core.author = "Proyecto El Teorema del Sí"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
